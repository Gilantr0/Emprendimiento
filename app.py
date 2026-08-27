import os
import shutil
import re
import subprocess
from datetime import datetime
from fastapi import FastAPI, HTTPException
from fastapi.responses import HTMLResponse, FileResponse
from pydantic import BaseModel
from pypdf import PdfReader
import docx
from PIL import Image

# 1. Declaración obligatoria de la instancia de FastAPI (DEBE ir antes de los decoradores @app)
app = FastAPI(title="BrainDoc Local Manager - AI Metadata Edition")

# Configuración de categorías y reglas de autotagging
CATEGORIES = {
    "Documentos": [".pdf", ".docx", ".txt", ".xlsx", ".pptx", ".csv"],
    "Imágenes": [".jpg", ".jpeg", ".png", ".gif", ".svg", ".webp"],
    "Audio_Video": [".mp3", ".wav", ".mp4", ".mkv", ".mov"],
    "Código": [".py", ".js", ".html", ".css", ".json", ".sh"],
    "Archivos_Comprimidos": [".zip", ".tar", ".gz", ".7z", ".rar"]
}

TAG_RULES = {
    "Facturación/Finanzas": ["factura", "cuenta", "recibo", "pago", "invoice", "rut", "dian", "impuesto", "total", "saldo"],
    "Redes/IT": ["cisco", "router", "switch", "vlan", "etherchannel", "ip", "subred", "linux", "ssh", "python"],
    "Contratos/Legal": ["contrato", "acuerdo", "cláusula", "firma", "legal", "certificación", "acta"],
    "Proyectos/Documentación": ["especificación", "requerimiento", "braindoc", "propuesta", "manual", "informe"]
}

class PathRequest(BaseModel):
    directory: str

class FilePathRequest(BaseModel):
    path: str

def get_category(extension):
    ext = extension.lower()
    for category, extensions in CATEGORIES.items():
        if ext in extensions:
            return category
    return "Otros"

def extract_text_and_metadata(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    text = ""
    metadata = {}

    try:
        if ext == ".txt":
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read(2000)
        elif ext == ".pdf":
            reader = PdfReader(filepath)
            metadata["pages"] = len(reader.pages)
            for page in reader.pages[:3]:
                text += page.extract_text() or ""
        elif ext == ".docx":
            doc = docx.Document(filepath)
            metadata["paragraphs"] = len(doc.paragraphs)
            text = " ".join([p.text for p in doc.paragraphs[:15]])
        elif ext in [".jpg", ".jpeg", ".png"]:
            with Image.open(filepath) as img:
                metadata["resolution"] = f"{img.width}x{img.height}"
                metadata["format"] = img.format
    except Exception as e:
        metadata["error"] = f"No se pudo extraer contenido: {str(e)}"

    extracted_tags = []
    text_lower = text.lower()
    for tag, keywords in TAG_RULES.items():
        if any(re.search(r'\b' + re.escape(kw) + r'\b', text_lower) for kw in keywords):
            extracted_tags.append(tag)

    if not extracted_tags and text_lower:
        extracted_tags.append("General/Texto")

    return extracted_tags, metadata, text[:150]

# 2. Rutas de la Aplicación (@app)

@app.get("/logo.jpeg")
def get_logo():
    # Busca la imagen en la misma carpeta del script
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Busca posibles nombres de archivo para el logo
    for filename in ["logo.jpeg", "logo.jpg", "logo.png", "WhatsApp Image 2026-08-12 at 19.27.38.jpeg"]:
        logo_path = os.path.join(base_dir, filename)
        if os.path.exists(logo_path):
            return FileResponse(logo_path)
            
    raise HTTPException(status_code=404, detail="Imagen del logo no encontrada en el directorio.")

@app.post("/api/scan")
def scan_directory(req: PathRequest):
    if not os.path.exists(req.directory):
        raise HTTPException(status_code=404, detail="Ruta no encontrada")
    
    files_info = []
    for entry in os.scandir(req.directory):
        if entry.is_file() and not entry.name.startswith('.'):
            stats = entry.stat()
            mod_time = datetime.fromtimestamp(stats.st_mtime)
            ext = os.path.splitext(entry.name)[1]
            
            tags, metadata, snippet = extract_text_and_metadata(entry.path)

            files_info.append({
                "name": entry.name,
                "path": entry.path,
                "size_kb": round(stats.st_size / 1024, 2),
                "modified": mod_time.strftime("%Y-%m-%d %H:%M:%S"),
                "category": get_category(ext),
                "tags": tags,
                "metadata": metadata,
                "snippet": snippet
            })
    return {"directory": req.directory, "files": files_info}

@app.post("/api/open")
def open_file(req: FilePathRequest):
    if not os.path.exists(req.path):
        raise HTTPException(status_code=404, detail="El archivo ya no existe en esa ruta")
    if not os.path.isfile(req.path):
        raise HTTPException(status_code=400, detail="La ruta no corresponde a un archivo")

    try:
        # os.startfile abre el archivo con la aplicación predeterminada de Windows
        # (ej: PDF con el lector de PDF, .docx con Word, imágenes con el visor de fotos)
        os.startfile(req.path)
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"No se pudo abrir el archivo: {str(e)}")

@app.post("/api/reveal")
def reveal_file(req: FilePathRequest):
    if not os.path.exists(req.path):
        raise HTTPException(status_code=404, detail="El archivo ya no existe en esa ruta")

    try:
        # Abre el Explorador de Windows en la carpeta del archivo, con el archivo ya seleccionado
        subprocess.run(f'explorer /select,"{os.path.normpath(req.path)}"')
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"No se pudo abrir el Explorador: {str(e)}")

@app.get("/", response_class=HTMLResponse)
def index():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    html_path = os.path.join(base_dir, "index.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return f.read()

if __name__ == "__main__":
    import uvicorn
    # host="0.0.0.0" hace que el servidor escuche en todas las interfaces de red,
    # no solo en localhost, para que otros equipos de la misma red puedan conectarse
    uvicorn.run(app, host="0.0.0.0", port=8000)